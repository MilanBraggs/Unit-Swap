import streamlit as st
from docx import Document
import re
from io import BytesIO


# Conversion function                 
def convert_units(number, unit):
    conversion_factors = {
        'cells/µL': (1000000, 'cells/L'),
        'mg/kg': (0.000001, ''),
        'µg/kg':(0.000000001,''),
        'µg/kg/minute': (0.000000001, '/minute'),
        'U/kg': (0.001, 'U/g'),
        'mg': (0.001, 'g'),
        '°C': (lambda x: x * 9/5 + 32, '°F'),
        'ng/ml': (0.001, 'mg/L'),
        'µg/ml': (1.0, 'mg/L'),
        'µg': (0.001, 'mg')
    }
    
    if unit in conversion_factors:
        factor, new_unit = conversion_factors[unit]
        if callable(factor):
            converted_number = factor(number)
        else:
            converted_number = number * factor
        return converted_number, new_unit
    return number, unit

# Function to replace text in a paragraph
def replace_text_in_paragraph(paragraph, pattern): 
    full_text = "".join([run.text for run in paragraph.runs])
    matches = list(re.finditer(pattern, full_text))
    new_text = full_text

    for match in reversed(matches):  # Reverse to handle replacements correctly
        original_text = match.group(0)
        number = float(match.group(1))
        unit = match.group(3)
        converted_number, converted_unit = convert_units(number, unit)
        converted_text = f"{converted_number:.2e} {converted_unit}"  
        new_text = new_text[:match.start()] + converted_text + new_text[match.end():]

    # Update runs
    if paragraph.runs:
        paragraph.clear()
        paragraph.add_run(new_text)


# Streamlit app
st.title("Unit Conversion in Word Document")

# Upload file section
with st.sidebar:
    st.header("Options")
    uploaded_file = st.file_uploader("Choose a Word document", type=["docx"])
    
if uploaded_file is not None:
    try:
        # Read the uploaded Word document
        doc = Document(uploaded_file)
        
        # Define the pattern for unit conversion
        pattern = r'\b(\d+(\.\d+)?)\s?(cells/µL|mg/kg|µg/kg|µg/kg/minute|U/kg|mg|µg/ml|ng/ml|µg|°C)\b' 
        
        # Make a copy of the original document for displaying original content
        original_doc = Document()
        for para in doc.paragraphs:
            original_doc.add_paragraph(para.text)

        # Process each paragraph and replace text
        for para in doc.paragraphs:
            replace_text_in_paragraph(para, pattern)
        
        # Save the processed document to a BytesIO object
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Display the original and converted documents side by side
        with st.container():
            col1, col2 = st.columns(2)

            with col1:
                st.header("Original Document")
                for para in original_doc.paragraphs:
                    st.write(para.text)

            with col2:
                st.header("Converted Document")
                for para in doc.paragraphs:
                    st.write(para.text)
        
        # Download button for the processed document
        st.download_button(
            label="Download Processed Document",
            data=buffer,
            file_name='processed_document.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        
    except Exception as e:
        st.error(f"Error: {e}")
else:
    st.info("Please upload a Word document to proceed.")

                    
            

